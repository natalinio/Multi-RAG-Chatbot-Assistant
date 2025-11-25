"""
Optimized document processor with larger semantic chunks for better RAG performance.
Target: 2000-4000 character chunks to preserve complete JSON configurations.
"""

import json
import logging
from pathlib import Path
from docx import Document
from datetime import datetime
import hashlib
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class OptimizedDocumentProcessor:
    """
    Process ETL documentation with optimized chunk sizes for RAG.
    Target: 2000-4000 characters to keep JSON templates complete.
    """
    
    def __init__(self):
        self.data_dir = Path(__file__).parent
        self.examples_dir = self.data_dir / "examples"
        self.output_dir = self.data_dir / "processed"
        self.output_dir.mkdir(exist_ok=True)
        
        # Optimal chunk sizes for RAG
        self.min_chunk_size = 2000  # Minimum chars to keep context
        self.max_chunk_size = 4000  # Maximum to avoid token limits
        self.target_chunk_size = 3000  # Ideal size
        
    def process_docx_file(self, doc_path: Path) -> list:
        """Process a single DOCX file into optimized chunks."""
        logger.info(f"Processing DOCX: {doc_path.name}...")
        
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")
        
        doc = Document(doc_path)
        chunks = []
        
        # Extract all text with structure preservation
        current_section = ""
        current_subsection = ""
        current_content = ""
        chunk_id = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if not text:
                continue
            
            # Detect major sections (numbered 1., 2., 3., etc.)
            major_section_match = re.match(r'^(\d+)\.\s+(.+)$', text)
            if major_section_match and len(text) < 100:
                # Save previous section if substantial
                if len(current_content) > 500:
                    chunks.append(self._create_chunk(
                        chunk_id, current_section, current_subsection, current_content
                    ))
                    chunk_id += 1
                
                current_section = text
                current_subsection = ""
                current_content = text + "\n\n"
                continue
            
            # Detect subsections
            is_subsection = (
                para.style.name.startswith('Heading') or
                text.endswith(':') or
                (text.isupper() and len(text) < 80)
            )
            
            if is_subsection and len(current_content) > self.max_chunk_size:
                # Chunk is too large, split here
                chunks.append(self._create_chunk(
                    chunk_id, current_section, current_subsection, current_content
                ))
                chunk_id += 1
                current_subsection = text
                current_content = text + "\n\n"
            elif is_subsection:
                current_subsection = text
                current_content += text + "\n\n"
            else:
                current_content += text + "\n"
            
            # Split if exceeded max size at natural boundary
            if len(current_content) > self.max_chunk_size:
                chunks.append(self._create_chunk(
                    chunk_id, current_section, current_subsection, current_content
                ))
                chunk_id += 1
                current_content = ""
        
        # Save last chunk
        if len(current_content) > 500:
            chunks.append(self._create_chunk(
                chunk_id, current_section, current_subsection, current_content
            ))
        
        # Process tables separately (they contain important structured data)
        table_chunks = self._extract_tables(doc, chunk_id)
        chunks.extend(table_chunks)
        
        logger.info(f"Created {len(chunks)} optimized chunks")
        return chunks
    
    def _create_chunk(self, chunk_id: int, section: str, subsection: str, content: str) -> dict:
        """Create a chunk with metadata."""
        content = content.strip()
        
        # Determine content type
        content_type = 'documentation'
        if 'json' in content.lower() or '{' in content:
            content_type = 'configuration_example'
        elif 'table' in section.lower() or subsection.lower() in ['tables']:
            content_type = 'table'
        
        # Extract keywords
        keywords = self._extract_keywords(content)
        
        # Generate semantic title from section and subsection
        title_parts = []
        if section and section != 'General':
            title_parts.append(section)
        if subsection:
            title_parts.append(subsection)
        semantic_title = " - ".join(title_parts) if title_parts else "ETL Configuration Guide"
        
        return {
            'id': f'opt_chunk_{chunk_id}_{hashlib.md5(content.encode()).hexdigest()[:8]}',
            'title': semantic_title,  # NEW: Semantic-friendly title
            'section': section or 'General',
            'subsection': subsection or '',
            'content': content,
            'content_type': content_type,
            'keywords': keywords,
            'word_count': len(content.split()),
            'character_count': len(content),
            'source_document': 'ETL_Configuration.docx',
            'processed_date': datetime.now().isoformat(),
            'metadata': {
                'chunk_strategy': 'optimized_semantic',
                'target_size': self.target_chunk_size
            }
        }
    
    def _extract_tables(self, doc: Document, start_id: int) -> list:
        """Extract tables as separate chunks."""
        table_chunks = []
        chunk_id = start_id
        
        for table_idx, table in enumerate(doc.tables):
            # Convert table to markdown-like format
            table_content = f"## Table {table_idx + 1}\n\n"
            
            # Get headers
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                table_content += "| " + " | ".join(headers) + " |\n"
                table_content += "|" + "|".join(['---' for _ in headers]) + "|\n"
                
                # Get data rows
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_content += "| " + " | ".join(cells) + " |\n"
            
            # Only keep substantial tables
            if len(table_content) > 200:
                # Extract table topic from first column header or content
                table_title = f"Table {table_idx + 1}"
                if table.rows and table.rows[0].cells:
                    first_header = table.rows[0].cells[0].text.strip()
                    if first_header:
                        table_title = f"Reference Table - {first_header}"
                
                table_chunks.append({
                    'id': f'opt_chunk_table_{chunk_id}_{hashlib.md5(table_content.encode()).hexdigest()[:8]}',
                    'title': table_title,  # NEW: Semantic title
                    'section': 'Tables',
                    'subsection': f'Table {table_idx + 1}',
                    'content': table_content,
                    'content_type': 'table',
                    'keywords': self._extract_keywords(table_content),
                    'word_count': len(table_content.split()),
                    'character_count': len(table_content),
                    'source_document': 'ETL_Configuration.docx',
                    'processed_date': datetime.now().isoformat(),
                    'metadata': {
                        'chunk_strategy': 'table_extraction',
                        'table_index': table_idx
                    }
                })
                chunk_id += 1
        
        logger.info(f"Extracted {len(table_chunks)} table chunks")
        return table_chunks
    
    def process_json_file(self, json_path: Path) -> dict:
        """Process a single JSON configuration file."""
        logger.info(f"Processing JSON: {json_path.name}...")
        
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            
            # Extract metadata
            domain = config.get('domain', 'Unknown')
            entity = config.get('entity', json_path.stem)
            layer = config.get('layer', 'Unknown')
            market = config.get('market', 'Unknown')
            process_type = config.get('process_requested', 'Unknown')
            
            # Create formatted content
            content_parts = [
                f"## ETL Configuration Example: {entity}",
                f"\n**Domain:** {domain}",
                f"**Layer:** {layer}",
                f"**Process Type:** {process_type}",
                f"**Market:** {market}",
                "\n### Dependencies"
            ]
            
            # Add dependencies
            inbound = config.get('dependencyInbound', [])
            outbound = config.get('dependencyOutbound', [])
            content_parts.append(f"- **Inbound:** {', '.join(inbound) if inbound else 'None'}")
            content_parts.append(f"- **Outbound:** {', '.join(outbound) if outbound else 'None'}")
            
            # Add configuration structure overview
            content_parts.append("\n### Configuration Structure\n")
            
            # Identify main sections
            for key in config.keys():
                if key.startswith(('I0_', 'I1_', 'I2_', 'D0_', 'D1_', 'D2_', 'D3_', 'IF0_', 'IF1_', 'IF2_', 'IF3_', 'R0_', 'R1_')):
                    section_type = config[key].get('type', 'N/A')
                    content_parts.append(f"**{key}:**")
                    content_parts.append(f"- Type: {section_type}")
                    
                    # Add key properties
                    section_data = config[key]
                    for prop_key in ['use-case', 'prcs-name', 'table_name', 'source', 'sink', 'target']:
                        if prop_key in section_data:
                            prop_value = section_data[prop_key]
                            if isinstance(prop_value, dict):
                                content_parts.append(f"- {prop_key}: {prop_value.get('type', 'complex object')}")
                            else:
                                content_parts.append(f"- {prop_key}: {prop_value}")
                    content_parts.append("")
            
            # Add complete JSON
            content_parts.append("\n### Complete JSON")
            content_parts.append("```json")
            content_parts.append(json.dumps(config, indent=2, ensure_ascii=False))
            content_parts.append("```\n")
            
            content = "\n".join(content_parts)
            
            # Generate semantic title for better search ranking
            process_stages = {
                'I0_': 'Data Ingestion', 'I1_': 'Data Extraction', 'I2_': 'Data Transformation',
                'D0_': 'Data Processing', 'D1_': 'Data Transformation', 'D2_': 'Data Loading', 'D3_': 'Data Loading',
                'IF0_': 'Interface', 'IF1_': 'Interface', 'IF2_': 'Interface', 'IF3_': 'Interface',
                'R0_': 'Reporting', 'R1_': 'Reporting'
            }
            
            # Detect main process type from configuration
            main_stage = 'Configuration'
            for key in config.keys():
                for prefix, stage in process_stages.items():
                    if key.startswith(prefix):
                        main_stage = stage
                        break
                if main_stage != 'Configuration':
                    break
            
            # Build semantic title: "[SourceType] [Stage] - [Layer] Layer"
            source_type = 'Generic'
            if 'I1_data_extract_process' in config and 'source' in config['I1_data_extract_process']:
                source_type = config['I1_data_extract_process']['source'].get('type', 'Generic')
            elif 'D1_data_extract_process' in config and 'source' in config['D1_data_extract_process']:
                source_type = config['D1_data_extract_process']['source'].get('type', 'Generic')
            
            semantic_title = f"{source_type} {main_stage} - {layer} Layer"
            if market and market != 'GLB':
                semantic_title += f" ({market})"
            
            return {
                "id": hashlib.md5(entity.encode()).hexdigest(),
                "title": semantic_title,  # NEW: Semantic-friendly title
                "content": content,
                "metadata": {
                    "type": "configuration_example",
                    "entity": entity,
                    "domain": domain,
                    "layer": layer,
                    "market": market,
                    "process_type": process_type,
                    "source_file": json_path.name,
                    "has_dependencies": len(inbound) > 0
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing JSON {json_path.name}: {e}")
            return None
    
    def process_all_documents(self) -> list:
        """Process all DOCX and JSON files in examples directory."""
        all_chunks = []
        
        if not self.examples_dir.exists():
            raise FileNotFoundError(f"Examples directory not found: {self.examples_dir}")
        
        # Process all DOCX files
        docx_files = list(self.examples_dir.glob("*.docx"))
        logger.info(f"Found {len(docx_files)} DOCX files to process")
        
        for docx_file in docx_files:
            try:
                chunks = self.process_docx_file(docx_file)
                all_chunks.extend(chunks)
                logger.info(f"✅ {docx_file.name}: {len(chunks)} chunks")
            except Exception as e:
                logger.error(f"❌ Failed to process {docx_file.name}: {e}")
        
        # Process all JSON files
        json_files = list(self.examples_dir.glob("*.json"))
        logger.info(f"Found {len(json_files)} JSON files to process")
        
        for json_file in json_files:
            try:
                chunk = self.process_json_file(json_file)
                if chunk:
                    all_chunks.append(chunk)
                    logger.info(f"✅ {json_file.name}: 1 chunk")
            except Exception as e:
                logger.error(f"❌ Failed to process {json_file.name}: {e}")
        
        logger.info(f"Total chunks from all files: {len(all_chunks)}")
        return all_chunks
    
    def _extract_keywords(self, text: str) -> list:
        """Extract ETL-related keywords."""
        keywords = []
        keyword_list = [
            'bronze', 'silver', 'gold', 'ingestion', 'transformation', 'validation',
            'sapcdc', 'sapbw', 'sap', 'adls', 'delta', 'parquet', 'json', 'azure',
            'i0', 'i1', 'i2', 'i3', 'd0', 'd1', 'd2', 'd3', 'd4',
            't0', 't1', 't2', 'r0', 'r1', 'if0', 'if1', 'if2', 'if3',
            'source', 'sink', 'destination', 'configuration', 'template',
            'master', 'common', 'domain', 'entity', 'layer', 'process',
            'extract', 'load', 'transform', 'integrate', 'refresh',
            'salesforce', 'profisee', 'nielsen', 'sql', 'databricks'
        ]
        
        text_lower = text.lower()
        for keyword in keyword_list:
            if keyword in text_lower:
                keywords.append(keyword)
        
        return list(set(keywords))
    
    def save_chunks(self, chunks: list):
        """Save optimized chunks to JSON."""
        output_file = self.output_dir / "processed_content.json"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(chunks, f, indent=2, ensure_ascii=False)
        
        # Create summary - handle different chunk structures
        sizes = []
        for c in chunks:
            if 'character_count' in c:
                sizes.append(c['character_count'])
            elif 'content' in c:
                sizes.append(len(c['content']))
        
        # Get content types with safe access
        content_types = {}
        for c in chunks:
            ctype = c.get('content_type') or c.get('metadata', {}).get('type', 'unknown')
            content_types[ctype] = content_types.get(ctype, 0) + 1
        
        # Calculate word counts
        total_words = 0
        for c in chunks:
            if 'word_count' in c:
                total_words += c['word_count']
            elif 'content' in c:
                total_words += len(c['content'].split())
        
        summary = {
            'total_chunks': len(chunks),
            'chunk_strategy': 'optimized_semantic_with_json',
            'target_chunk_size': self.target_chunk_size,
            'chunk_size_stats': {
                'min': min(sizes) if sizes else 0,
                'max': max(sizes) if sizes else 0,
                'avg': sum(sizes) // len(sizes) if sizes else 0,
                'median': sorted(sizes)[len(sizes)//2] if sizes else 0
            },
            'content_types': content_types,
            'total_words': total_words,
            'total_characters': sum(sizes),
            'processed_date': datetime.now().isoformat(),
            'source_files': 'all_examples_directory'
        }
        
        summary_file = self.output_dir / "processing_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved {len(chunks)} chunks to {output_file}")
        logger.info(f"Summary: {summary}")
        
        return summary


def main():
    """Main processing function."""
    try:
        processor = OptimizedDocumentProcessor()
        chunks = processor.process_all_documents()
        summary = processor.save_chunks(chunks)
        
        print("\n" + "="*70)
        print("OPTIMIZED DOCUMENT PROCESSING COMPLETED")
        print("="*70)
        print(f"Total chunks: {summary['total_chunks']}")
        print(f"Target chunk size: {summary['target_chunk_size']} characters")
        print(f"\nChunk size statistics:")
        print(f"  Minimum: {summary['chunk_size_stats']['min']:,} chars")
        print(f"  Maximum: {summary['chunk_size_stats']['max']:,} chars")
        print(f"  Average: {summary['chunk_size_stats']['avg']:,} chars")
        print(f"  Median: {summary['chunk_size_stats']['median']:,} chars")
        print(f"\nContent types:")
        for ctype, count in summary['content_types'].items():
            print(f"  {ctype}: {count}")
        print(f"\nTotal characters: {summary['total_characters']:,}")
        print(f"Total words: {summary['total_words']:,}")
        print("="*70)
        print("\n✅ Documents processed with optimized chunk sizes for RAG!")
        print("   Next step: Run upload_to_search.py to reindex")
        
    except Exception as e:
        logger.error(f"Processing failed: {str(e)}")
        import traceback
        traceback.print_exc()
        raise


if __name__ == "__main__":
    main()
